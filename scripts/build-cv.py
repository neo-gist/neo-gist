#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build-cv.py — 홈페이지 데이터로 CV(DOCX) 표를 자동으로 채운다.
  template : files/cv/cv_template.docx  (양식 마스터 — 사람이 서식/정적 섹션 관리)
  sources  : data/publications/ref.bib · patents.json ,  data/members/professor/talks.json · cv.json
  output   : files/cv/donghokang.docx   (표만 최신 데이터로 교체)
  PDF 변환 : 별도(LibreOffice) — GitHub Action 에서 수행
규칙:
  · Selected Publications = ref.bib 의 selected={true} (홈페이지 Highlights 와 동일)
  · Main / Contributing = 강동호가 1저자·공동1저자·교신저자면 Main
  · review=true → 제목 앞 (Invited Review)
  · 커버는 표시하지 않음
실행:  python3 scripts/build-cv.py
"""
import re, json, sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

KO_FONT = "NanumGothic"     # 한국어(East Asian) 글꼴
JCR_BLUE = RGBColor(0x11, 0x55, 0xCC)  # JCR 상위 10% 저널 표시 색

ROOT = Path(__file__).resolve().parent.parent
PROF = ROOT / "data/members/professor"
PUB  = ROOT / "data/publications"
CVDIR = ROOT / "files/cv"
FONT, SIZE = "Times New Roman", Pt(10)

def load_json(p): return json.loads(Path(p).read_text(encoding="utf-8"))

# ---------- BibTeX ----------
def bib_clean(s):
    s = s or ""
    s = s.replace(r"$\alpha$", "α")
    s = re.sub(r"\$([^$]*)\$", r"\1", s)
    s = s.replace("{", "").replace("}", "")
    return re.sub(r"\s+", " ", s).strip()

def parse_bib(text):
    blocks, cur = [], None
    for line in text.splitlines():
        if re.match(r"^\s*@\w+\s*\{", line):
            if cur is not None: blocks.append(cur)
            cur = line + "\n"
        elif cur is not None:
            cur += line + "\n"
            if re.match(r"^\s*\}\s*$", line):
                blocks.append(cur); cur = None
    if cur: blocks.append(cur)
    out = []
    for i, e in enumerate(blocks):
        f = {}
        for m in re.finditer(r"(\w[\w+]*)\s*=\s*\{([\s\S]*?)\}\s*,?\s*(?=\n\s*\w[\w+]*\s*=|\n\s*\}\s*$)", e):
            f[m.group(1).lower()] = m.group(2)
        authors = [a.strip() for a in re.split(r"\s+and\s+", f.get("author", "")) if a.strip()]
        corr, joint = set(), set()
        for t in (f.get("author+an", "") or "").split(";"):
            mm = re.search(r"(\d+)\s*=\s*([a-z]+)", t, re.I)
            if mm:
                n, r = int(mm.group(1)), mm.group(2).lower()
                if r.startswith("correspond"): corr.add(n)
                if r.startswith("joint"): joint.add(n)
        out.append(dict(
            i=i, title=bib_clean(f.get("title")), journal=bib_clean(f.get("journal")),
            year=(f.get("year") or "").strip(), authors=authors, corr=corr, joint=joint,
            review=bool(re.search(r"review", f.get("type", ""), re.I)),
            selected=bool(re.search(r"true|1|yes", f.get("selected", ""), re.I)),
        ))
    return out

def norm(s): return re.sub(r"[^a-z]", "", (s or "").lower())
def split_name(raw):
    c = re.sub(r"[{}]", "", raw); c = re.sub(r"\s+", " ", c).strip()
    if "," in c:
        last, rest = c.split(",", 1); return last.strip(), rest.strip()
    parts = c.split(" "); return parts[-1], " ".join(parts[:-1])
def initials(given):
    toks = [t for t in given.split(" ") if t]
    return " ".join("-".join((p[0].upper() + ".") for p in tok.split("-") if p) for tok in toks)

def kang_index(p):
    for idx, raw in enumerate(p["authors"]):
        last, given = split_name(raw)
        if norm(last) == "kang" and norm(given).startswith("dongho"): return idx
    return -1
def is_main(p):
    k = kang_index(p)
    return k >= 0 and (k == 0 or (k+1) in p["joint"] or (k+1) in p["corr"])

# 저자 셀 run 스펙: [(text, dict(bold,italic,underline,super)), ...]
def author_runs(p):
    specs = []
    for idx, raw in enumerate(p["authors"]):
        n = idx + 1
        last, given = split_name(raw)
        name = (initials(given) + " " + last).strip()
        is_kang = (norm(last) == "kang" and norm(given).startswith("dongho"))
        if idx: specs.append((", ", {}))
        specs.append((name, {"bold": is_kang, "underline": is_kang}))
        if n in p["joint"]: specs.append(("+", {"super": True, "bold": is_kang, "underline": is_kang}))
        if n in p["corr"]:  specs.append(("*", {"bold": is_kang, "underline": is_kang}))
    return specs

# ---------- DOCX 헬퍼 ----------
def set_run(r, spec):
    r.font.name = FONT; r.font.size = SIZE
    if spec.get("bold"): r.bold = True
    if spec.get("italic"): r.italic = True
    if spec.get("underline"): r.underline = True
    if spec.get("super"): r.font.superscript = True
    if spec.get("color"): r.font.color.rgb = spec["color"]

def fill_cell(cell, specs, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # 기존 문단 제거 후 새 문단 하나
    for ppar in list(cell.paragraphs):
        ppar._element.getparent().remove(ppar._element)
    para = cell.add_paragraph(); para.alignment = align
    para.paragraph_format.space_before = Pt(0); para.paragraph_format.space_after = Pt(0)
    last = None
    for text, spec in specs:
        if text == "\n":
            if last is None: last = para.add_run("")
            last.add_break(); continue
        r = para.add_run(text); set_run(r, spec)
        last = r
    return para

def clear_data_rows(table):
    for row in list(table.rows)[1:]:
        row._element.getparent().remove(row._element)

def norm_journal(s):
    return re.sub(r"[^a-z0-9]", "", (s or "").lower().replace("&", "and"))

def build_jcr(raw):
    return {yr: {norm_journal(k): v for k, v in tbl.items()} for yr, tbl in (raw or {}).items()}

def jcr_top10(p, jcr):
    """게재연도 −1년 JCR 기준 상위 10% 이내면 True (사이트와 동일 규칙)."""
    if not p["year"].isdigit():
        return False
    tbl = jcr.get(str(int(p["year"]) - 1))
    if not tbl:
        return False
    rec = tbl.get(norm_journal(p["journal"]))
    try:
        return rec is not None and float(rec.get("pct", 999)) <= 10
    except (TypeError, ValueError):
        return False

def journal_cell(p, jmap, top10=False):
    ab = jmap.get(p["journal"], p["journal"])
    jspec = {"italic": True, "bold": True, "color": JCR_BLUE} if top10 else {"italic": True}
    return [(ab, jspec), ("\n", {}), ("(" + p["year"] + ")", {})]

def title_specs(p):
    if p["review"]:
        return [("(Invited Review) ", {"italic": True}), (p["title"], {})]
    return [(p["title"], {})]

def fill_pub_table(table, papers, jmap, jcr):
    clear_data_rows(table)
    n = len(papers)
    for idx, p in enumerate(papers):
        no = n - idx  # 최신이 큰 번호
        row = table.add_row()
        fill_cell(row.cells[0], [(str(no), {"bold": True})])
        fill_cell(row.cells[1], title_specs(p), WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(row.cells[2], journal_cell(p, jmap, jcr_top10(p, jcr)))
        fill_cell(row.cells[3], author_runs(p), WD_ALIGN_PARAGRAPH.CENTER)

def patent_country(pt):
    s = (pt.get("app") or pt.get("reg") or "").upper()
    if re.search(r"\bUS\b|USA|^US", s): return "USA"
    if re.search(r"^KR|^10-|KOREA", s): return "Korea"
    return "USA" if pt.get("region") == "intl" else "Korea"

def kang_bold_specs(text):
    """강동호(영문 Dong-Ho Kang / 국문 강동호)를 굵게."""
    specs, last = [], 0
    for m in re.finditer(r"Dong-Ho Kang|강동호", text):
        if m.start() > last: specs.append((text[last:m.start()], {}))
        specs.append((m.group(0), {"bold": True})); last = m.end()
    if last < len(text): specs.append((text[last:], {}))
    return specs or [(text, {})]

def fill_patents(table, patents):
    clear_data_rows(table)
    pats = sorted(patents, key=lambda x: str(x.get("y", "")))  # 오래된 것부터
    for idx, pt in enumerate(pats):
        row = table.add_row()
        fill_cell(row.cells[0], [(str(idx + 1), {"bold": True})])
        fill_cell(row.cells[1], [(pt.get("t", ""), {})])
        fill_cell(row.cells[2], [(patent_country(pt), {"italic": True}), ("\n", {}), ("(" + str(pt.get("y", "")) + ")", {})])
        fill_cell(row.cells[3], kang_bold_specs(pt.get("a", "")))

def fill_talks(table, talks):
    clear_data_rows(table)
    ts = sorted(talks, key=lambda x: str(x.get("y", "")), reverse=True)  # 최신 먼저
    n = len(ts)
    for idx, t in enumerate(ts):
        row = table.add_row()
        fill_cell(row.cells[0], [(str(n - idx), {"bold": True})])
        title = ([("(Invited) ", {"italic": True})] if re.search(r"invit", t.get("type", ""), re.I) else []) + [(t.get("t", ""), {})]
        fill_cell(row.cells[1], title, WD_ALIGN_PARAGRAPH.CENTER)
        conf = [(t.get("conf", ""), {"italic": True})]
        if t.get("place"): conf += [("\n", {}), (t["place"], {})]
        if t.get("date"): conf += [("\n", {}), (t["date"], {})]
        fill_cell(row.cells[2], conf)
        fill_cell(row.cells[3], [("D.-H. Kang", {"bold": True, "underline": True}), ("*", {})])

def set_col_widths(table, widths_in):
    """열 폭 고정(inch). LibreOffice 가 존중하도록 fixed layout + 모든 셀 폭 지정."""
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout'); tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    # tblGrid 갱신
    grid = table._tbl.find(qn('w:tblGrid'))
    if grid is not None:
        cols = grid.findall(qn('w:gridCol'))
        for c, w in zip(cols, widths_in):
            c.set(qn('w:w'), str(int(w * 1440)))
    for row in table.rows:
        for i, w in enumerate(widths_in):
            row.cells[i].width = Inches(w)

def header_height(table):
    tr = table.rows[0]._tr; trPr = tr.find(qn('w:trPr'))
    if trPr is None: return None
    h = trPr.find(qn('w:trHeight'))
    return (h.get(qn('w:val')), h.get(qn('w:hRule'))) if h is not None else None

def set_header_height(table, val, rule=None):
    tr = table.rows[0]._tr; trPr = tr.get_or_add_trPr()
    h = trPr.find(qn('w:trHeight'))
    if h is None:
        h = OxmlElement('w:trHeight'); trPr.append(h)
    h.set(qn('w:val'), str(val))
    if rule:
        h.set(qn('w:hRule'), rule)
    elif h.get(qn('w:hRule')) is not None:
        del h.attrib[qn('w:hRule')]

def rows_no_split(table, header_repeat=True):
    """행이 페이지 경계에서 잘리지 않도록(cantSplit) + 표가 다음 페이지로 넘어가면 헤더 반복."""
    for i, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        trPr.append(OxmlElement('w:cantSplit'))
        if header_repeat and i == 0:
            trPr.append(OxmlElement('w:tblHeader'))

def set_all_eastasia(doc, font=KO_FONT):
    """문서 전체 run 의 East Asian(한글) 글꼴을 지정."""
    def fix(runs):
        for r in runs:
            rpr = r._element.get_or_add_rPr()
            rfonts = rpr.find(qn('w:rFonts'))
            if rfonts is None:
                rfonts = OxmlElement('w:rFonts'); rpr.insert(0, rfonts)
            rfonts.set(qn('w:eastAsia'), font)
    for p in doc.paragraphs:
        fix(p.runs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    fix(p.runs)

def trim_tail(doc):
    """마지막 표 뒤에 남는 빈 문단들을 정리 — 불필요한 빈 페이지 방지."""
    from docx.text.paragraph import Paragraph
    body = doc.element.body
    kids = list(body)
    tbl_idx = [i for i, e in enumerate(kids) if e.tag == qn('w:tbl')]
    if not tbl_idx:
        return
    tail = [e for e in kids[tbl_idx[-1] + 1:] if e.tag == qn('w:p')]
    for e in tail[1:]:
        body.remove(e)                      # 표 뒤 첫 문단 하나만 남기고 제거
    if tail:
        p = Paragraph(tail[0], doc)
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        pf = p.paragraph_format
        pf.space_before = Pt(0); pf.space_after = Pt(0); pf.line_spacing = Pt(2)
        run = p.add_run('')
        run.font.size = Pt(1)               # 꼬리 문단 높이 최소화

def add_jcr_legend(doc, heading_text):
    """섹션 제목 오른쪽 끝에 작은 파란 범례를 추가."""
    sec = doc.sections[0]
    usable = sec.page_width - sec.left_margin - sec.right_margin
    for para in doc.paragraphs:
        if para.style and para.style.name == "Heading 1" and para.text.strip().lower() == heading_text.lower():
            para.paragraph_format.tab_stops.add_tab_stop(usable, WD_TAB_ALIGNMENT.RIGHT)
            r = para.add_run("\tTop 10% Journal (JCR)")
            r.font.size = Pt(8); r.font.bold = True; r.font.small_caps = False
            r.font.color.rgb = JCR_BLUE
            return

def J(s):
    return (s or "").replace("&amp;", "&")

def expand_abbr(s):
    """CV 에서는 약어 대신 풀네임 사용."""
    s = re.sub(r"\bEECS\b", "Electrical Engineering and Computer Science", s)
    s = re.sub(r"\bEEE\b", "Electrical and Electronic Engineering", s)
    return s

def _find_h1(doc, text):
    for p in doc.paragraphs:
        if p.style and p.style.name == "Heading 1" and p.text.strip().lower() == text.lower():
            return p
    return None

def clear_section_body(doc, title):
    """제목 다음부터 다음 Heading1 직전까지의 문단 삭제 → 삽입 앵커(다음 heading Paragraph) 반환."""
    h = _find_h1(doc, title)
    if h is None:
        return None
    heads = [p._element for p in doc.paragraphs if p.style and p.style.name == "Heading 1"]
    h_el = h._element
    nxt = heads[heads.index(h_el) + 1] if heads.index(h_el) + 1 < len(heads) else None
    sib = h_el.getnext()
    while sib is not None and sib is not nxt:
        nx = sib.getnext()
        if sib.tag == qn('w:p'):
            sib.getparent().remove(sib)
        else:
            break
        sib = nx
    return Paragraph(nxt, doc) if nxt is not None else None

def build_static_sections(doc, profile, service, engagement, cv):
    """헤더·Education·Appointments·Awards 를 데이터로 다시 채우고,
       Professional Service·Public Engagement 섹션을 추가한다."""
    sec = doc.sections[0]
    RIGHT = sec.page_width - sec.left_margin - sec.right_margin
    EDU_COL, APPT_COL = Inches(0.95), Inches(1.95)

    def mk(anchor):
        p = anchor.insert_paragraph_before()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        return p
    def R(p, text, bold=False, italic=False):
        if not text: return
        run = p.add_run(text); run.bold = bold; run.italic = italic; run.font.name = FONT
    def ltab(p, pos): p.paragraph_format.tab_stops.add_tab_stop(pos, WD_TAB_ALIGNMENT.LEFT)
    def rtab(p): p.paragraph_format.tab_stops.add_tab_stop(RIGHT, WD_TAB_ALIGNMENT.RIGHT)

    # 이름
    for p in doc.paragraphs:
        if p.style and p.style.name == "Name":
            for run in list(p.runs): run.text = ""
            (p.runs[0] if p.runs else p.add_run("")).text = J(profile.get("en", ""))
            break

    # 헤더 표(연락처)
    contact = profile.get("contact", {})
    cell = doc.tables[0].rows[0].cells[0]
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    def hline(text):
        p = cell.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        R(p, text); return p
    hline("")
    hline(J(profile.get("role", "")))
    for dl in profile.get("dept", []): hline(J(dl))
    hline("")
    if contact.get("tel"): hline("Tel: " + contact["tel"])
    if cv.get("mobile"): hline("Mobile: " + cv["mobile"])
    if contact.get("email"): hline("Mailing address: " + contact["email"])

    # Education
    a = clear_section_body(doc, "Education")
    if a is not None:
        for e in profile.get("education", []):
            p = mk(a); ltab(p, EDU_COL); rtab(p)
            R(p, J(e.get("deg", "")), bold=True); R(p, "\t"); R(p, J(e.get("org", ""))); R(p, "\t"); R(p, str(e.get("y", "")))
            p2 = mk(a); p2.paragraph_format.left_indent = EDU_COL; R(p2, J(e.get("major", "")))
            note = J(e.get("note", ""))
            if note:
                p3 = mk(a); p3.paragraph_format.left_indent = EDU_COL
                if "(" in note:
                    x, y = note.split("(", 1); R(p3, x); R(p3, "(" + y, italic=True)
                else:
                    R(p3, note)
            mk(a)

    # Appointments — 역할+기간을 한 줄, 소속·부서는 아래 줄(들여쓰기)로 (기간 줄바꿈 방지)
    a = clear_section_body(doc, "Appointments")
    if a is not None:
        for x in profile.get("experience", []):
            p = mk(a); rtab(p)
            R(p, J(x.get("role", "")), bold=True); R(p, "\t"); R(p, J(x.get("period", "")))
            p2 = mk(a); p2.paragraph_format.left_indent = Inches(0.3); R(p2, expand_abbr(J(x.get("org", ""))))
            if x.get("unit"):
                p3 = mk(a); p3.paragraph_format.left_indent = Inches(0.3); R(p3, expand_abbr(J(x.get("unit", ""))))
            mk(a)

    # Awards
    a = clear_section_body(doc, "Awards")
    if a is not None:
        for aw in service.get("awards", []):
            p = mk(a); rtab(p)
            R(p, J(aw.get("t", "")), bold=True); R(p, "\t"); R(p, str(aw.get("y", "")))
            R(mk(a), J(aw.get("org", "")))
            if aw.get("note"): R(mk(a), J(aw["note"]), italic=True)
            mk(a)

def rename_heading(doc, old, new):
    for para in doc.paragraphs:
        if para.style and para.style.name == "Heading 1" and para.text.strip().lower() == old.lower():
            for r in para.runs: r.text = ""
            (para.runs[0] if para.runs else para.add_run("")).text = new
            return True
    return False

def main():
    cv = load_json(PROF / "cv.json"); jmap = cv.get("journal_abbr", {})
    jcr = build_jcr(load_json(PUB / "jcr.json"))
    engagement = load_json(PROF / "engagement.json")
    profile = load_json(PROF / "profile.json")
    service = load_json(PROF / "service.json")
    papers = [p for p in parse_bib((PUB / "ref.bib").read_text(encoding="utf-8"))
              if re.match(r"^\d{4}$", p["year"]) and p["title"]]
    ydesc = lambda p: (int(p["year"]), -p["i"])
    selected = sorted([p for p in papers if p["selected"]], key=ydesc, reverse=True)
    main_a   = sorted([p for p in papers if is_main(p)], key=ydesc, reverse=True)
    contrib  = sorted([p for p in papers if not is_main(p)], key=ydesc, reverse=True)
    patents = load_json(PUB / "patents.json").get("patents", [])
    talks   = load_json(PROF / "talks.json")

    doc = Document(str(CVDIR / "cv_template.docx"))
    rename_heading(doc, "5 Representative publications", "Selected Publications")
    rename_heading(doc, "Presentations in International Conferences", "Invited Talks")

    # 정적 섹션(소속·학력·경력·수상) 자동 생성 + Professional Service·Public Engagement 추가
    build_static_sections(doc, profile, service, engagement, cv)

    tables = doc.tables  # 0=header,1=Selected,2=Main,3=Contributing,4=Patents,5=Presentations(Invited Talks)
    fill_pub_table(tables[1], selected, jmap, jcr)
    fill_pub_table(tables[2], main_a, jmap, jcr)
    fill_pub_table(tables[3], contrib, jmap, jcr)
    fill_patents(tables[4], patents)
    fill_talks(tables[5], talks)
    # JCR 상위 10% 범례를 논문 표 제목 오른쪽에
    for h in ("Selected Publications", "Full-publication list (main-authored)", "Full-publication list (Contributing-authored)"):
        add_jcr_legend(doc, h)
    # 논문·특허 표는 모두 동일한 고정 열 폭으로 통일  [No, Title, Journal, Authors]
    PUB_W = [0.39, 2.26, 1.08, 2.56]
    for ti in (1, 2, 3, 4):
        set_col_widths(tables[ti], PUB_W)
    # Invited Talks: 바깥 폭·No·Title 은 동일하게 맞추고 Conference 만 넓게 / Authors 좁게
    set_col_widths(tables[5], [0.39, 2.26, 2.70, 0.95])
    # 헤더 행 높이를 Selected 표 기준으로 전부 통일
    ref = header_height(tables[1])
    if ref:
        for ti in (2, 3, 4, 5):
            set_header_height(tables[ti], ref[0], ref[1])
    # 모든 표: 행이 페이지 경계에서 잘리지 않게 + 헤더 반복
    for ti in range(1, 6):
        rows_no_split(tables[ti])
    # 모든 한국어를 나눔고딕으로
    set_all_eastasia(doc)
    # 마지막 표 뒤 빈 문단 정리(빈 페이지 방지)
    trim_tail(doc)

    out = CVDIR / "donghokang.docx"
    doc.save(str(out))
    print(f"DOCX 생성: {out}")
    print(f"  Selected {len(selected)} · Main {len(main_a)} · Contributing {len(contrib)} · Patents {len(patents)} · Talks {len(talks)}")

if __name__ == "__main__":
    main()
